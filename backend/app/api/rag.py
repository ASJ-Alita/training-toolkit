"""
RAG 知识库问答 API
"""

import os
import json
import shutil
import time
import requests
import numpy as np
from fastapi import APIRouter, HTTPException, UploadFile, File
from sqlalchemy.orm import Session

from app.core.database import SessionLocal
from app.models.models_ext import Document, ChatMessage
from app.models.schemas_ext import (
    ChatRequest, ChatResponse, LLMConfigRequest,
    LLM_PROVIDERS, SILICONFLOW_EMBED,
    CHUNK_SIZE, CHUNK_OVERLAP, TOP_K, SIMILARITY_THRESHOLD,
)

router = APIRouter(prefix="/api/rag", tags=["RAG 知识库"])

# ── 数据目录 ──────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DOCS_DIR = os.path.join(BASE_DIR, "data", "documents")
INDEX_DIR = os.path.join(BASE_DIR, "data", "indexes")
CONFIG_FILE = os.path.join(BASE_DIR, "data", "rag_config.json")

os.makedirs(DOCS_DIR, exist_ok=True)
os.makedirs(INDEX_DIR, exist_ok=True)


# ── 配置管理 ──────────────────────────────────────────────────────────────

def _load_config() -> dict:
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"llm_provider": "siliconflow", "api_key": ""}


def _save_config(config: dict):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


@router.get("/config")
def get_rag_config():
    config = _load_config()
    return {
        "llm_provider": config.get("llm_provider", "siliconflow"),
        "llm_providers": {k: v["name"] for k, v in LLM_PROVIDERS.items()},
        "api_key_set": bool(config.get("api_key", "")),
    }


@router.post("/config")
def update_rag_config(data: LLMConfigRequest):
    config = _load_config()
    config["llm_provider"] = data.llm_provider
    if data.api_key:
        config["api_key"] = data.api_key
    _save_config(config)
    return {"message": "配置已保存"}


# ── 文档处理 ──────────────────────────────────────────────────────────────

def _extract_text(filepath: str) -> str:
    """从文件中提取文本（支持 txt/pdf/docx）"""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".txt":
        for enc in ["utf-8", "gbk", "gb2312"]:
            try:
                with open(filepath, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    elif ext == ".pdf":
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
            return text.strip()
        except ImportError:
            raise HTTPException(500, "需要安装 pdfplumber：pip install pdfplumber")

    elif ext in [".docx", ".doc"]:
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(500, "需要安装 python-docx：pip install python-docx")
    else:
        raise HTTPException(400, f"不支持的文件格式：{ext}")


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list:
    """文本分块（滑动窗口）"""
    if len(text) <= chunk_size:
        return [{"text": text, "chunk_id": 0}]
    chunks = []
    start = 0
    chunk_id = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if end < len(text):
            for sep in ["。", "！", "？", ".\n", "?\n", "!\n", "\n"]:
                last_sep = chunk.rfind(sep)
                if last_sep > chunk_size * 0.6:
                    end = start + last_sep + len(sep)
                    chunk = text[start:end]
                    break
        chunks.append({"text": chunk.strip(), "chunk_id": chunk_id})
        chunk_id += 1
        start = end - overlap
    return chunks


# ── 向量存储（轻量 JSON + numpy，无需 FAISS 依赖） ───────────────────────

def _get_index_path(provider: str) -> tuple:
    """获取索引文件路径"""
    index_file = os.path.join(INDEX_DIR, f"vectors_{provider}.npy")
    meta_file = os.path.join(INDEX_DIR, f"chunks_{provider}.json")
    return index_file, meta_file


def _load_vectors(provider: str) -> tuple:
    """加载向量数据"""
    index_file, meta_file = _get_index_path(provider)
    if os.path.exists(index_file) and os.path.exists(meta_file):
        vectors = np.load(index_file)
        with open(meta_file, "r", encoding="utf-8") as f:
            chunks = json.load(f)
        return vectors, chunks
    return np.empty((0, SILICONFLOW_EMBED["dimension"]), dtype=np.float32), []


def _save_vectors(provider: str, vectors: np.ndarray, chunks: list):
    """保存向量数据"""
    index_file, meta_file = _get_index_path(provider)
    np.save(index_file, vectors)
    with open(meta_file, "w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)


def _get_embeddings(texts: list, provider: str, api_key: str) -> list:
    """获取文本向量"""
    if not api_key:
        raise HTTPException(400, "请先配置 API Key")

    if provider in ("siliconflow", "zhihui") or not provider:
        url = SILICONFLOW_EMBED["url"]
        model = SILICONFLOW_EMBED["name"]
    elif provider == "openai":
        url = "https://api.openai.com/v1/embeddings"
        model = "text-embedding-3-small"
    else:
        url = SILICONFLOW_EMBED["url"]
        model = SILICONFLOW_EMBED["name"]

    all_vectors = []
    batch_size = 50
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        resp = requests.post(url, headers=headers,
                             json={"model": model, "input": batch}, timeout=60)
        resp.raise_for_status()
        for item in resp.json()["data"]:
            all_vectors.append(item["embedding"])

    return all_vectors


# ── 文档上传 API ─────────────────────────────────────────────────────────

@router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """上传文档到知识库"""
    # 检查格式
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".txt", ".pdf", ".docx", ".doc"]:
        raise HTTPException(400, "仅支持 .txt / .pdf / .docx 格式")

    config = _load_config()
    if not config.get("api_key"):
        raise HTTPException(400, "请先配置 API Key（RAG 向量化需要）")

    # 保存文件
    save_path = os.path.join(DOCS_DIR, file.filename)
    if os.path.exists(save_path):
        base, ex = os.path.splitext(file.filename)
        save_path = os.path.join(DOCS_DIR, f"{base}_{int(time.time())}{ex}")

    with open(save_path, "wb") as f:
        f.write(await file.read())

    # 提取文本
    text = _extract_text(save_path)
    if not text.strip():
        os.remove(save_path)
        raise HTTPException(400, "文档内容为空")

    # 分块
    chunks = _chunk_text(text)
    texts = [c["text"] for c in chunks]

    # 向量化
    provider = config.get("llm_provider", "siliconflow")
    try:
        embeddings = _get_embeddings(texts, provider, config["api_key"])
    except Exception as e:
        os.remove(save_path)
        raise HTTPException(500, f"向量化失败：{str(e)}")

    # 存入向量索引
    vectors, existing_chunks = _load_vectors(provider)
    new_metas = []
    for i, c in enumerate(chunks):
        new_metas.append({
            "text": c["text"],
            "source": file.filename,
            "chunk_id": c["chunk_id"],
            "ingest_time": time.strftime("%Y-%m-%d %H:%M:%S"),
        })

    if vectors.size > 0:
        new_vecs = np.array(embeddings, dtype=np.float32)
        vectors = np.vstack([vectors, new_vecs])
    else:
        vectors = np.array(embeddings, dtype=np.float32)
    existing_chunks.extend(new_metas)
    _save_vectors(provider, vectors, existing_chunks)

    # 记录到数据库
    db = SessionLocal()
    try:
        doc = Document(
            filename=file.filename,
            file_path=save_path,
            file_size=os.path.getsize(save_path),
            file_type=ext,
            chunks_count=len(chunks),
            status="completed",
        )
        db.add(doc)
        db.commit()
        doc_id = doc.id
    finally:
        db.close()

    return {"id": doc_id, "filename": file.filename, "chunks": len(chunks), "status": "completed"}


@router.get("/documents")
def list_documents():
    """获取已上传文档列表"""
    db = SessionLocal()
    try:
        docs = db.query(Document).order_by(Document.created_at.desc()).all()
        return [{
            "id": d.id, "filename": d.filename, "file_size": d.file_size,
            "chunks_count": d.chunks_count, "status": d.status,
            "created_at": d.created_at,
        } for d in docs]
    finally:
        db.close()


@router.get("/stats")
def get_kb_stats():
    """获取知识库统计"""
    config = _load_config()
    provider = config.get("llm_provider", "siliconflow")
    _, chunks = _load_vectors(provider)
    sources = list(set(c.get("source", "unknown") for c in chunks))

    db = SessionLocal()
    try:
        doc_count = db.query(Document).filter(Document.status == "completed").count()
    finally:
        db.close()

    return {
        "total_documents": doc_count,
        "total_chunks": len(chunks),
        "sources": sources,
        "llm_provider": provider,
        "dimension": SILICONFLOW_EMBED["dimension"],
    }


@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int):
    """删除文档（从 DB 删除，向量暂不清理）"""
    db = SessionLocal()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            raise HTTPException(404, "文档不存在")
        if os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        db.delete(doc)
        db.commit()
        return {"message": "已删除"}
    finally:
        db.close()


@router.delete("/clear")
def clear_knowledge_base():
    """清空知识库"""
    # 清理数据库
    db = SessionLocal()
    try:
        db.query(ChatMessage).delete()
        db.query(Document).delete()
        db.commit()
    finally:
        db.close()

    # 清理向量索引
    for fname in os.listdir(INDEX_DIR):
        fpath = os.path.join(INDEX_DIR, fname)
        if os.path.isfile(fpath):
            os.remove(fpath)

    # 清理文档
    for fname in os.listdir(DOCS_DIR):
        fpath = os.path.join(DOCS_DIR, fname)
        if os.path.isfile(fpath):
            os.remove(fpath)

    return {"message": "知识库已清空"}


# ── 问答 API ──────────────────────────────────────────────────────────────

@router.post("/chat")
def chat(data: ChatRequest):
    """RAG 智能问答"""
    if not data.question.strip():
        raise HTTPException(400, "问题不能为空")

    config = _load_config()
    if not config.get("api_key"):
        raise HTTPException(400, "请先配置 API Key")

    provider = config.get("llm_provider", "siliconflow")

    # 1. 问题向量化
    try:
        q_vectors = _get_embeddings([data.question], provider, config["api_key"])
        q_embedding = np.array(q_vectors[0], dtype=np.float32)
    except Exception as e:
        raise HTTPException(500, f"向量化失败：{str(e)}")

    # 2. 检索
    vectors, chunks = _load_vectors(provider)
    if vectors.size == 0 or not chunks:
        return ChatResponse(
            answer="知识库为空，请先上传文档后再提问。",
            sources=[], question=data.question,
        )

    # 归一化并计算余弦相似度
    q_norm = q_embedding / (np.linalg.norm(q_embedding) + 1e-8)
    v_norms = vectors / (np.linalg.norm(vectors, axis=1, keepdims=True) + 1e-8)
    similarities = (v_norms @ q_norm)

    # Top-K
    top_k = min(TOP_K, len(chunks))
    top_indices = np.argsort(similarities)[::-1][:top_k]

    results = []
    for idx in top_indices:
        score = float(similarities[idx])
        if score >= SIMILARITY_THRESHOLD:
            results.append((chunks[idx], score))

    if not results:
        return ChatResponse(
            answer="在知识库中未找到与您问题相关的内容。\n\n可能的原因：\n• 知识库为空（请先上传文档）\n• 问题与文档内容不相关",
            sources=[], question=data.question,
        )

    # 3. 构建上下文
    context_parts = []
    sources_list = []
    for i, (chunk, score) in enumerate(results, 1):
        context_parts.append(f"【来源{i}】{chunk['source']}（相关度：{score:.2f}）\n{chunk['text']}")
        sources_list.append({"text": chunk["text"][:200] + "...", "source": chunk["source"], "score": round(score, 3)})
    context = "\n\n".join(context_parts)

    # 4. 调用 LLM
    system_prompt = (
        "你是一位专业的知识库问答助手，名为「智库助手」。\n"
        "你的职责是：基于提供的知识库内容，准确回答用户的问题。\n\n"
        "重要规则：\n"
        "1. 只根据【参考资料】中的内容回答，不要编造信息\n"
        "2. 如果知识库内容不足以完整回答，请明确说明\n"
        "3. 回答要专业、准确、简洁\n"
        "4. 可以引用【来源】标注内容出处\n"
        "5. 用中文回答"
    )

    user_prompt = (
        f"【用户问题】\n{data.question}\n\n"
        f"【参考资料】\n{context}\n\n"
        f"请根据参考资料回答用户问题。如果资料不足，请说明。"
    )

    llm_cfg = LLM_PROVIDERS.get(provider, LLM_PROVIDERS["siliconflow"])
    try:
        resp = requests.post(
            llm_cfg["chat_url"],
            headers={"Authorization": f"Bearer {config['api_key']}", "Content-Type": "application/json"},
            json={
                "model": llm_cfg["model"],
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": llm_cfg["temperature"],
                "max_tokens": llm_cfg["max_tokens"],
            },
            timeout=60,
        )
        resp.raise_for_status()
        answer = resp.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        raise HTTPException(500, f"LLM 调用失败：{str(e)}")

    # 5. 保存对话历史
    db = SessionLocal()
    try:
        db.add(ChatMessage(session_id=data.session_id, role="user", content=data.question))
        db.add(ChatMessage(session_id=data.session_id, role="assistant",
                           content=answer, sources=json.dumps(sources_list, ensure_ascii=False)))
        db.commit()
    finally:
        db.close()

    return ChatResponse(answer=answer, sources=sources_list, question=data.question)


@router.get("/chat/history/{session_id}")
def get_chat_history(session_id: str):
    """获取对话历史"""
    db = SessionLocal()
    try:
        msgs = db.query(ChatMessage).filter(
            ChatMessage.session_id == session_id
        ).order_by(ChatMessage.id).all()
        return [{
            "role": m.role,
            "content": m.content,
            "sources": json.loads(m.sources) if m.sources else [],
        } for m in msgs]
    finally:
        db.close()


@router.delete("/chat/history/{session_id}")
def clear_chat_history(session_id: str):
    """清空对话历史"""
    db = SessionLocal()
    try:
        db.query(ChatMessage).filter(ChatMessage.session_id == session_id).delete()
        db.commit()
        return {"message": "对话历史已清空"}
    finally:
        db.close()
